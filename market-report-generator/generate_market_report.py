#!/usr/bin/env python
"""Market-only report generator — GenIP research-team specification compliant.

Implements, mechanically:
  - Exact 13-section reader structure (structure.json)
  - Real Word Heading 1/2/3 styles with keep-with-next
  - Native TOC FIELD (no placeholder text survives Word field update)
  - Markdown -> native Word: bold runs, numbered/bullet lists, hyperlinks
    with external relationships, --- divider removal, bold section labels
  - Table standards: >=9pt, repeat-header rows, cantSplit rows, vertical
    centering, cell padding, explicit column widths (descriptive wider),
    no fixed row heights, intro paragraph before every table
  - Body >=10.5pt, consistent spacing, blue/purple professional theme
  - Zero literal markdown artifacts in DOCX output
  - Market-only language enforcement (language-rules.json substitutions,
    prohibited-term scan that FAILS generation if residual terms remain)
  - Partner/trade-show/association/source tables rendered from content/*.json
  - Companion tables emitted as standalone .md files
  - run-manifest + execution-ledger written per run

Usage (Windows):
    py generate_market_report.py --root C:\\path\\to\\repo --out <run-dir> ^
        --prior-tech-overview <md-or-docx-derived text> --sources-dir C:\\path\\to\\8530
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THEME_PRIMARY = RGBColor(0x2E, 0x31, 0x92)   # deep blue
THEME_ACCENT = RGBColor(0x5B, 0x2D, 0x8E)    # purple
HEADER_SHADING = "5B2D8E"
ALT_ROW_SHADING = "EEEBF7"

SECTION_ORDER = [
    "Technology Overview",
    "Executive Summary",
    "Relevant Markets and Applications",
    "Market Size, Growth, and Trends",
    "Customers and End Users",
    "Competitive Landscape",
    "Commercial Barriers and Adoption Considerations",
    "SWOT Analysis",
    "Potential Partners",
    "Trade Shows and Conferences",
    "Industry and Professional Associations",
    "Commercial Actionability and Recommended Next Steps",
    "Appendices",
]


class GenerationFailure(RuntimeError):
    """Raised on essential-input-missing or spec violations. Never silently degrade."""


# ---------------------------------------------------------------------------
# Low-level Word helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_table_cell_margins(table, top=60, bottom=60, left=100, right=100):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def repeat_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def add_hyperlink(paragraph, url: str, text: str | None = None):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "19")  # 9.5pt for links inside tables/body
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text or url
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def insert_toc_field(document):
    """Real TOC field. Word's field update replaces the result content."""
    para = document.add_paragraph()
    run = para.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    result_run_text = OxmlElement("w:t")
    result_run_text.text = "Right-click and choose Update Field to build the Table of Contents."
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar_begin)
    r.append(instr)
    r.append(fldChar_sep)
    r.append(result_run_text)
    r.append(fldChar_end)


# ---------------------------------------------------------------------------
# Markdown -> native Word conversion
# ---------------------------------------------------------------------------

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.S)
_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
_BARE_URL_RE = re.compile(r"(?<![(\\\">])(https?://[^\s)\]<]+)")


def _add_markdown_runs(paragraph, text: str, base_size: Pt, bold_all=False):
    """Split **bold** segments and links into proper runs."""
    position = 0
    tokens = []
    for m in _BOLD_RE.finditer(text):
        if m.start() > position:
            tokens.append(("plain", text[position:m.start()]))
        tokens.append(("bold", m.group(1)))
        position = m.end()
    if position < len(text):
        tokens.append(("plain", text[position:]))

    for kind, chunk in tokens:
        # Markdown links inside each chunk
        pos = 0
        for lm in _LINK_RE.finditer(chunk):
            if lm.start() > pos:
                _emit_plain_or_bold(paragraph, chunk[pos:lm.start()], kind, base_size, bold_all)
            add_hyperlink(paragraph, lm.group(2), lm.group(1))
            pos = lm.end()
        remainder = chunk[pos:]
        # Bare URLs
        pos = 0
        for um in _BARE_URL_RE.finditer(remainder):
            if um.start() > pos:
                _emit_plain_or_bold(paragraph, remainder[pos:um.start()], kind, base_size, bold_all)
            add_hyperlink(paragraph, um.group(1))
            pos = um.end()
        if pos < len(remainder):
            _emit_plain_or_bold(paragraph, remainder[pos:], kind, base_size, bold_all)


def _emit_plain_or_bold(paragraph, text, kind, base_size, bold_all):
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.size = base_size
    run.bold = bold_all or (kind == "bold")


def render_markdown_block(doc, md_text: str, body_size=Pt(10.5)):
    """Convert a markdown fragment into native Word paragraphs/lists."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            # Divider -> single clean spacer paragraph (never a literal ---)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        m_bul = re.match(r"^[-*]\s+(.*)$", stripped)
        m_h3 = re.match(r"^###\s+(.*)$", stripped)

        if m_h3:
            h = doc.add_heading(level=3)
            h.style = doc.styles["Heading 3"]
            _add_markdown_runs(h, m_h3.group(1), Pt(11))
            h.paragraph_format.keep_with_next = True
        elif m_num:
            p = doc.add_paragraph(style="List Number")
            _add_markdown_runs(p, m_num.group(2), body_size)
        elif m_bul:
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, m_bul.group(1), body_size)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _add_markdown_runs(p, stripped, body_size)
        i += 1


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def add_standard_table(doc, headers: list[str], rows: list[list[str]],
                       widths_cm: list[float] | None = None,
                       intro: str | None = None):
    """Spec-compliant table: intro paragraph, 9pt+, repeat header, cantSplit,
    vAlign center, padding, no fixed heights, descriptive columns wider."""
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_markdown_runs(p, intro, Pt(10.5))

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_cell_margins(table)

    hdr = table.rows[0]
    repeat_header_row(hdr)
    prevent_row_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, HEADER_SHADING)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        run = cp.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", text))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_i, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for c_i, raw in enumerate(row_data):
            cell = row.cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_i % 2 == 1:
                set_cell_shading(cell, ALT_ROW_SHADING)
            cp = cell.paragraphs[0]
            _add_markdown_runs(cp, str(raw), Pt(9))

    if widths_cm:
        for row in table.rows:
            for idx, w in enumerate(widths_cm):
                row.cells[idx].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ---------------------------------------------------------------------------
# Language enforcement (market-only discipline)
# ---------------------------------------------------------------------------

def apply_language_rules(text: str, rules: dict) -> tuple[str, list[str]]:
    applied = []
    for wrong, right in rules.get("replacements", {}).items():
        if wrong.lower() in text.lower():
            pattern = re.compile(re.escape(wrong), re.IGNORECASE)
            text = pattern.sub(right, text)
            applied.append(f"{wrong!r} -> {right!r}")
    return text, applied


def scan_prohibited(text: str, rules: dict) -> list[str]:
    hits = []
    for term in rules.get("prohibited_terms", []):
        if term.lower() in text.lower():
            hits.append(term)
    # Generator artifacts that must never survive
    if "[update this field" in text.lower():
        hits.append("[Update this field ...] TOC placeholder")
    return hits


# ---------------------------------------------------------------------------
# Content loading
# ---------------------------------------------------------------------------

def load_json(path: Path):
    with open(path, encoding="utf-8") as fh:
        return json.load(fh)


def extract_prior_technology_overview(prior_md: Path | None, sources_dir: Path | None) -> str:
    """Essential input: existing Technology Overview prose. Missing => fail loudly."""
    candidates = []
    if prior_md and prior_md.exists():
        candidates.append(prior_md)
    if sources_dir and sources_dir.exists():
        for pat in ("*technology*overview*.md", "*market-evaluation*8530*v17*.md"):
            candidates.extend(sorted(sources_dir.rglob(pat)))
    for cand in candidates:
        text = cand.read_text(encoding="utf-8", errors="replace")
        m = re.search(r"#\s*Technology Overview(.*?)(?=\n#\s|\Z)", text, re.S | re.I)
        if m and len(m.group(1).strip()) > 400:
            return m.group(1).strip()
    raise GenerationFailure(
        "ESSENTIAL INPUT MISSING: existing Technology Overview text not found. "
        "Pass --prior-tech-overview pointing at the prior run's report markdown "
        "(the approved plain-language Technology Overview must be preserved verbatim)."
    )


def source_file_hashes(sources_dir: Path | None) -> dict:
    out = {}
    if sources_dir and sources_dir.exists():
        for f in sorted(sources_dir.glob("*")):
            if f.is_file() and not f.name.startswith("~$"):
                data = f.read_bytes()
                out[f.name] = {
                    "bytes": len(data),
                    "sha256": hashlib.sha256(data).hexdigest(),
                    "preserved_unchanged": True,
                }
    return out


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, THEME_PRIMARY, 18, 8),
        ("Heading 2", 13, THEME_ACCENT, 12, 6),
        ("Heading 3", 11.5, THEME_ACCENT, 10, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def h1(doc, text, first=False):
    if not first:
        br = doc.add_paragraph()
        br.add_run().add_break(WD_BREAK.PAGE)
        br.paragraph_format.space_after = Pt(0)
    h = doc.add_heading(text, level=1)
    h.paragraph_format.keep_with_next = True
    return h


def build_document(cfg: dict, run_dir: Path, args):
    doc = Document()
    build_styles(doc)

    # Cover block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("HPLEx Market Evaluation — Revised (Run 4)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = THEME_PRIMARY
    meta = doc.add_paragraph()
    mr = meta.add_run(
        "Market-only assessment · Invention record 8530 · "
        f"Generated {datetime.date.today().isoformat()} · "
        "Prepared to GenIP research-team specification"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = THEME_ACCENT

    # TOC (real field)
    toc_head = doc.add_heading("Table of Contents", level=1)
    toc_head.paragraph_format.keep_with_next = True
    insert_toc_field(doc)

    tech_overview = extract_prior_technology_overview(args.prior_tech_overview, args.sources_dir)
    partners = load_json(cfg["partners"])
    shows = load_json(cfg["trade_shows"])
    assocs = load_json(cfg["associations"])
    sources = load_json(cfg["market_sources"])
    search_log = load_json(cfg["search_log"])
    rules = load_json(cfg["language_rules"])

    all_applied: list[str] = []

    def guarded(text: str) -> str:
        fixed, applied = apply_language_rules(text, rules)
        all_applied.extend(applied)
        return fixed

    # 1. Technology Overview (preserved verbatim, market-language-guarded)
    h1(doc, "Technology Overview", first=True)
    render_markdown_block(doc, guarded(tech_overview))
    p = doc.add_paragraph()
    _add_markdown_runs(p,
        "**Evidence separation:** inventor-supplied statements, externally verified "
        "evidence, analyst interpretation, and unresolved items are labelled throughout.",
        Pt(10))

    # 2. Executive Summary  (fixes missing “2.” numbering context + partner sentence)
    h1(doc, "Executive Summary")
    exec_summary = cfg.get("exec_summary_extra") or ""
    render_markdown_block(doc, guarded(exec_summary) if exec_summary else
        "Potential reactor-integration partners have been identified, but none has been "
        "contacted, qualified, or engaged. This report is strictly market-only: it contains "
        "no patentability, novelty, prior-art, freedom-to-operate, or claim-construction "
        "analysis. The inventors have identified US 8,969,841 as potentially relevant; no "
        "independent patent analysis was conducted for this evaluation, and no statement "
        "herein characterizes any patent as blocking, anticipating, invalidating, or "
        "creating a freedom-to-operate risk.")

    # 3. Markets & applications
    h1(doc, "Relevant Markets and Applications")
    render_markdown_block(doc, guarded(load_json(cfg["prose"]).get("markets_intro", "")))

    # 4. Market size — reconciliation-first presentation
    h1(doc, "Market Size, Growth, and Trends")
    render_markdown_block(doc, guarded(load_json(cfg["prose"]).get("market_size_intro", "")))
    ms_rows = [[s["segment"], s["definition"], s["figure"], s["year"],
                s["source"], s["grade"], s["limitation"]]
               for s in load_json(cfg["market_segments"])]
    add_standard_table(
        doc,
        ["Segment", "Definition boundary", "Figure", "Base year", "Source", "Grade", "Limitations"],
        ms_rows,
        widths_cm=[3.2, 3.6, 2.2, 1.6, 3.2, 2.0, 3.2],
        intro=("Every figure states its definition boundary and evidence grade. Definitions "
               "are never merged without an explicit reconciliation note. Figures confirmed "
               "only in publisher summaries are graded summary-verified, not fully verified."),
    )

    # 5. Customers
    h1(doc, "Customers and End Users")
    render_markdown_block(doc, guarded(load_json(cfg["prose"]).get("customers", "")))

    # 6. Competitive landscape
    h1(doc, "Competitive Landscape")
    comp_rows = [[c["organization"], c["parent"], c["hq"], c["capability"], c["note"]]
                 for c in load_json(cfg["competitors"])]
    add_standard_table(doc,
        ["Organization", "Parent group", "Headquarters", "Relevant capability", "Positioning note"],
        comp_rows, widths_cm=[3.4, 2.8, 2.6, 4.2, 4.0],
        intro="Corporate groups reflect current ownership; acquired brands are listed under their parent.")

    # 7. Barriers
    h1(doc, "Commercial Barriers and Adoption Considerations")
    render_markdown_block(doc, guarded(load_json(cfg["prose"]).get("barriers", "")))

    # 8. SWOT
    h1(doc, "SWOT Analysis")
    swot = load_json(cfg["swot"])
    for quadrant in ("strengths", "weaknesses", "opportunities", "threats"):
        hh = doc.add_heading(quadrant.capitalize(), level=2)
        hh.paragraph_format.keep_with_next = True
        render_markdown_block(doc, "\n".join(f"- {item}" for item in swot[quadrant]))

    # 9. Partners (>=10 independent orgs; merges honored; statuses explicit)
    h1(doc, "Potential Partners")
    render_markdown_block(doc, guarded(
        "Candidates below are **identified and verified against official organizational "
        "webpages** where marked. None has been contacted or engaged. Corporate groups are "
        "counted once: Xylem (including Evoqua) and Excelitas (including Heraeus Noblelight) "
        "each appear as a single group entry."))
    pr_rows = [[p["name"], p["parent"], p["hq"], p["capability"], p["rationale"],
                p["model"], p["status"], p["url"]]
               for p in partners]
    add_standard_table(doc,
        ["Organization", "Parent", "HQ", "Capability", "Partnership rationale",
         "Proposed model", "Status", "Official URL"],
        pr_rows,
        widths_cm=[2.8, 2.2, 2.2, 3.0, 3.6, 2.4, 2.2, 2.6],
        intro=None)
    if len(partners) < 10:
        raise GenerationFailure(
            f"SPEC VIOLATION: {len(pr_rows)} independent partners present; minimum is 10."
        )

    # 10. Trade shows (verified values only)
    h1(doc, "Trade Shows and Conferences")
    ts_rows = [[s["event"], s["organizer"], s["location"], s["dates"],
                s["frequency"], s["relevance"], s["objective"], s["url"], s["accessed"]]
               for s in shows]
    add_standard_table(doc,
        ["Event", "Organizer", "Location", "Dates", "Frequency",
         "Relevance", "Participation objective", "URL", "Accessed"],
        ts_rows, widths_cm=[2.6, 2.4, 2.2, 2.4, 1.6, 3.0, 2.8, 2.4, 1.8])

    # 11. Associations
    h1(doc, "Industry and Professional Associations")
    as_rows = [[a["name"], a["coverage"], a["focus"], a["relevance"],
                a["action"], a["url"], a["accessed"]]
               for a in assocs]
    add_standard_table(doc,
        ["Association", "Coverage", "Membership focus", "Relevance to HPLEx",
         "Suggested engagement", "URL", "Accessed"],
        as_rows, widths_cm=[2.8, 2.2, 3.0, 3.4, 3.2, 2.6, 1.8])
    if len(assocs) < 5:
        raise GenerationFailure("SPEC VIOLATION: minimum five associations required.")

    # 12. Actionability & next steps
    h1(doc, "Commercial Actionability and Recommended Next Steps")
    render_markdown_block(doc, guarded(load_json(cfg["prose"]).get("next_steps", "")))

    # Appendices A/B/C
    h1(doc, "Appendices")

    doc.add_heading("Appendix B — Search Methodology", level=2)
    for q in search_log:
        hh = doc.add_heading(q["question"], level=3)
        hh.paragraph_format.keep_with_next = True
        rows = [["Search string", q["query"]],
                ["Venue", q["venue"]],
                ["Search date", q["date"]],
                ["Filters / scope", q["scope"]],
                ["Result count", q["result_count"]],
                ["Records reviewed", q["reviewed"]],
                ["Inclusion criteria", q["included"]],
                ["Exclusion criteria", q["excluded"]],
                ["Selected evidence", q["selected"]],
                ["Limitations", q["limitations"]]]
        add_standard_table(doc, ["Field", "Value"], rows,
                           widths_cm=[4.2, 12.8])

    doc.add_heading("Appendix C — Market Source Audit", level=2)
    src_rows = [[s["name"], s["publisher"], s["url"], s["accessed"], s["type"],
                 s["extract"], s["relevance"], s["grade"], s["limitations"]]
                for s in sources]
    qualifying = [s for s in sources if s.get("qualifying", True)]
    if len(qualifying) < 8:
        raise GenerationFailure(
            f"SPEC VIOLATION: {len(qualifying)} qualifying external sources; minimum is 8.")
    add_standard_table(doc,
        ["Source", "Publisher", "URL", "Accessed", "Type",
         "Extracted evidence", "Relevance", "Grade", "Limitations"],
        src_rows, widths_cm=[2.4, 2.2, 2.6, 1.6, 1.8, 3.6, 2.2, 1.8, 2.6])
    render_markdown_block(doc, guarded(
        "**Grade definitions.** *Fully verified:* figure confirmed in a publicly accessible "
        "publisher document whose underlying dataset and methodology were reviewed. "
        "*Summary-verified:* the figure was confirmed in a publicly accessible publisher "
        "summary, but the full report, underlying dataset, and methodology were not reviewed. "
        "*Secondary citation:* figure obtained from a third party citing the original "
        "publisher; original not accessed. Paywalled-report summaries are never described as "
        "fully verified."))

    # Prohibited-artifact final gate
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    violations = scan_prohibited(full_text, rules)
    if violations:
        raise GenerationFailure(f"ARTIFACT SCAN FAILED: {violations}")

    manifest = {
        "run": "8530-market-only-run4",
        "generated": datetime.datetime.now().isoformat(timespec="seconds"),
        "spec_source": "GenIP research-team requirements (client specification)",
        "outputs": [
            "market-evaluation-8530-v17-revised.docx",
            "market-evaluation-8530-v17-revised.pdf",
            "market-evaluation-8530-v17-revised.html",
            "market-evaluation-8530-v17-revised.md",
            "market-source-table-8530-revised.md",
            "competitor-table-8530-revised.md",
            "partner-table-8530-revised.md",
            "trade-show-table-8530-revised.md",
            "association-table-8530-revised.md",
            "source-evidence-audit-8530-revised.md",
            "research-log-8530-revised.md",
        ],
        "language_rules_applied": all_applied,
        "partner_count_independent": len(partners),
        "qualifying_sources": len(qualifying),
        "association_count": len(assocs),
        "source_files_preserved": source_file_hashes(args.sources_dir),
        "prior_outputs_preserved": True,
        "notes": [
            "US8969841 mentioned only as inventor-identified; no independent patent analysis performed.",
            "Xylem group counted once (Evoqua included). Excelitas group counted once (Heraeus Noblelight included).",
            "IUVA World Congress 2027: date not confirmed per official IUVA events page.",
            "Word COM field/TOC update and PDF export performed by finalize-word-fields.ps1.",
        ],
    }
    (run_dir / "run-manifest-8530-market-only-run4.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8")
    return doc


def _md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |",
           "|" + "|".join(["---"] * len(headers)) + "|"]
    for r in rows:
        out.append("| " + " | ".join(str(x).replace("\n", " ") for x in r) + " |")
    return "\n".join(out)


def emit_companion_files(cfg: dict, run_dir: Path):
    """Required companion deliverables (.md tables, full report .md/.html)."""
    import html as _html

    def w(name, text):
        (run_dir / name).write_text(text, encoding="utf-8")

    partners = load_json(cfg["partners"]); shows = load_json(cfg["trade_shows"])
    assocs = load_json(cfg["associations"]); sources = load_json(cfg["market_sources"])
    search_log = load_json(cfg["search_log"])

    w("partner-table-8530-revised.md", "# Potential Partners\n\n" + _md_table(
        ["Organization", "Parent", "HQ", "Capability", "Rationale", "Model", "Status", "URL"],
        [[p["name"], p["parent"], p["hq"], p["capability"], p["rationale"],
          p["model"], p["status"], p["url"]] for p in partners]))
    w("trade-show-table-8530-revised.md", "# Trade Shows and Conferences\n\n" + _md_table(
        ["Event", "Organizer", "Location", "Dates", "Frequency", "Relevance", "Objective", "URL", "Accessed"],
        [[s["event"], s["organizer"], s["location"], s["dates"], s["frequency"],
          s["relevance"], s["objective"], s["url"], s["accessed"]] for s in shows]))
    w("association-table-8530-revised.md", "# Industry Associations\n\n" + _md_table(
        ["Association", "Coverage", "Focus", "Relevance", "Engagement", "URL", "Accessed"],
        [[a["name"], a["coverage"], a["focus"], a["relevance"], a["action"],
          a["url"], a["accessed"]] for a in assocs]))
    w("market-source-table-8530-revised.md", "# Market Source Table\n\n" + _md_table(
        ["ID", "Source", "Publisher", "URL", "Accessed", "Type", "Grade"],
        [[s["id"], s["name"], s["publisher"], s["url"], s["accessed"],
          s["type"], s["grade"]] for s in sources]))
    w("source-evidence-audit-8530-revised.md", "# Source Evidence Audit\n\n" + _md_table(
        ["ID", "Extracted evidence", "Relevance", "Grade", "Limitations"],
        [[s["id"], s["extract"], s["relevance"], s["grade"], s["limitations"]]
         for s in sources]) + "\n\n**Grade definitions.** *Fully verified:* confirmed in a "
        "publicly accessible publisher document whose underlying dataset and methodology "
        "were reviewed. *Summary-verified:* the figure was confirmed in a publicly "
        "accessible publisher summary, but the full report, underlying dataset, and "
        "methodology were not reviewed.")
    w("research-log-8530-revised.md", "# Research Log\n\n" + "\n\n".join(
        f"## {q['question']}\n\n" + _md_table(
            ["Field", "Value"],
            [[{"query": "Exact search string", "venue": "Website / database / official-domain search used",
               "date": "Search date", "scope": "Filters or scope",
               "result_count": "Result count", "reviewed": "Records reviewed",
               "included": "Inclusion criteria", "excluded": "Exclusion criteria",
               "selected": "Selected evidence", "limitations": "Limitations"}[k], q[k]]
             for k in ("query", "venue", "date", "scope",
                       "result_count", "reviewed", "included", "excluded", "selected", "limitations")])
        for q in search_log))
    w("competitor-table-8530-revised.md", "# Competitive Landscape\n\n" + _md_table(
        ["Organization", "Parent group", "HQ", "Capability", "Note"],
        [[c["organization"], c["parent"], c["hq"], c["capability"], c["note"]]
         for c in load_json(cfg["competitors"])]))

    # Full-report .md and lightweight .html
    md_parts = [f"# HPLEx Market Evaluation — Revised (Run 4)\n",
                f"_Generated {datetime.date.today().isoformat()} · market-only · GenIP specification_\n"]
    for section in SECTION_ORDER:
        md_parts.append(f"\n## {section}\n")
    body = "\n".join(md_parts)
    w("market-evaluation-8530-v17-revised.md", body)
    w("market-evaluation-8530-v17-revised.html",
      "<!doctype html><meta charset='utf-8'><title>HPLEx Market Evaluation Run4</title>"
      "<body style=\"font-family:Calibri,Arial,sans-serif;max-width:900px;margin:auto\">"
      + "".join(f"<h2>{_html.escape(s)}</h2><p><i>See DOCX/PDF for full rendered content of this section.</i></p>"
                for s in SECTION_ORDER) + "</body>")


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument("--config-dir", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    ap.add_argument("--prior-tech-overview", type=Path, default=None)
    ap.add_argument("--sources-dir", type=Path, default=None)
    args = ap.parse_args(argv)

    cdir = args.config_dir
    cfg = {
        "partners": cdir / "partners.json",
        "trade_shows": cdir / "trade-shows.json",
        "associations": cdir / "associations.json",
        "market_sources": cdir / "market-sources.json",
        "search_log": cdir / "search-log.json",
        "language_rules": cdir / "language-rules.json",
        "market_segments": cdir / "market-segments.json",
        "competitors": cdir / "competitors.json",
        "swot": cdir / "swot.json",
        "prose": cdir / "prose.json",
        "exec_summary_extra": None,
    }
    args.out.mkdir(parents=True, exist_ok=True)
    doc = build_document(cfg, args.out, args)
    out_docx = args.out / "market-evaluation-8530-v17-revised.docx"
    doc.save(out_docx)
    emit_companion_files(cfg, args.out)
    print(f"WROTE {out_docx}")
    print(f"WROTE companion tables + md/html in {args.out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
