#!/usr/bin/env python
"""Executable acceptance contract — GenIP research-team specification.

Validates a generated run directory against the researcher's requirements as
machine-enforced FAIL-if rules. Exit code 0 = accepted; 1 = rejected with itemized
failures. The acceptance report lands in the run dir as evidence.

Usage:
    py validate_acceptance.py --run-dir <dir> [--language-rules <json>]
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


EXPECTED_FILES = [
    "market-evaluation-8530-v17-revised.docx",
    "market-evaluation-8530-v17-revised.html",
    "market-evaluation-8530-v17-revised.md",
    "market-source-table-8530-revised.md",
    "competitor-table-8530-revised.md",
    "partner-table-8530-revised.md",
    "trade-show-table-8530-revised.md",
    "association-table-8530-revised.md",
    "source-evidence-audit-8530-revised.md",
    "research-log-8530-revised.md",
    "execution-ledger-8530-market-only-run4.md",
    "run-manifest-8530-market-only-run4.json",
    "visual-qa-checklist-8530-run4.md",
]

REQUIRED_SECTION_ORDER = [
    "Technology Overview",
    "Executive Summary",
    "Relevant Markets and Applications",
    "Market Size, Growth, and Trends",
    "Customers and End Users",
    "Competitive Landscape",
    "Commercial Barriers and Adoption Considerations",
    "SWOT Analysis",
    "Potential Partners",
    "Trade Shows and Conferences",
    "Industry and Professional Associations",
    "Commercial Actionability and Recommended Next Steps",
    "Appendices",
]

DEFAULT_PROHIBITED = [
    "closest prior art",
    "potential ip obstacle",
    "file patent application",
    "file a patent application",
    "fto risk",
    "anticipates claim",
    "anticipated by",
    "obvious over",
]

_PDF_PENDING_NOTE = "word-finalization-pending.txt"


class AcceptanceReport:
    def __init__(self):
        self.checks: list[dict] = []

    def check(self, cid: str, passed: bool, detail: str = ""):
        self.checks.append({"id": cid, "pass": bool(passed), "detail": detail})

    @property
    def failures(self):
        return [c for c in self.checks if not c["pass"]]

    @property
    def passed(self):
        return not self.failures

    def to_dict(self):
        return {
            "accepted": self.passed,
            "total_checks": len(self.checks),
            "failed_checks": len(self.failures),
            "checks": self.checks,
        }


def _iter_table_text(doc):
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield cell.text


def _md_rows(path: Path) -> list[list[str]]:
    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if s.startswith("|") and "---" not in s:
            cells = [c.strip() for c in s.strip("|").split("|")]
            rows.append(cells)
    return rows[1:] if rows else []  # drop header


def validate(run_dir: Path, language_rules: Path | None = None) -> AcceptanceReport:
    rep = AcceptanceReport()
    docx_path = run_dir / "market-evaluation-8530-v17-revised.docx"

    # ---- Deliverable presence -------------------------------------------------
    missing = [f for f in EXPECTED_FILES if not (run_dir / f).exists()]
    rep.check("deliverables.all-present", not missing,
              f"missing: {missing}" if missing else f"{len(EXPECTED_FILES)} files present")
    rep.check("pdf.finalized-or-flagged",
              (run_dir / "market-evaluation-8530-v17-revised.pdf").exists()
              or (run_dir / _PDF_PENDING_NOTE).exists(),
              "PDF exported by Word COM, or pending-note written when Word unavailable")

    if not docx_path.exists():
        rep.check("docx.loadable", False, "DOCX missing; cannot run document checks")
        return rep

    try:
        doc = Document(docx_path)
        rep.check("docx.loadable", True)
    except Exception as exc:
        rep.check("docx.loadable", False, str(exc))
        return rep

    para_texts = [p.text for p in doc.paragraphs]
    full_text = "\n".join(para_texts + list(_iter_table_text(doc)))
    lower = full_text.lower()

    # ---- Domain 1: structure --------------------------------------------------
    h1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    pos = -1
    order_ok = True
    for req in REQUIRED_SECTION_ORDER:
        try:
            idx = h1_texts.index(req, pos + 1)
            pos = idx
        except ValueError:
            order_ok = False
            break
    rep.check("structure.section-order", order_ok,
              "13 required sections in exact order" if order_ok
              else f"sequence broken at '{req}' in {h1_texts}")
    rep.check("structure.exec-summary-heading", "Executive Summary" in h1_texts,
              "Executive Summary present as Heading 1")

    # ---- Domain 1/2: TOC field & placeholders ---------------------------------
    xml = doc.element.xml
    rep.check("toc.field-present", "TOC \\o" in xml or "TOC \\\\o" in xml or 'TOC \\o' in xml,
              "native TOC field instruction found in document XML")
    rep.check("toc.placeholder-absent", "[update this field" not in lower,
              "'[Update this field...]' must never survive")
    brackets = re.findall(r"\[[A-Z][A-Z0-9][A-Z0-9 _\-]{4,}\]", full_text)
    rep.check("artifacts.bracket-placeholders", not brackets,
              f"placeholders: {brackets[:5]}" if brackets else "none")

    # ---- Domain 2: markdown artifacts -----------------------------------------
    stars = full_text.count("**")
    rep.check("markdown.no-bold-markers", stars == 0,
              f"{stars} literal '**' occurrences" if stars else "zero '**' markers")
    dashes = [t for t in para_texts if t.strip() == "---"]
    rep.check("markdown.no-divider-paragraphs", not dashes,
              f"{len(dashes)} standalone '---' paragraphs")

    # ---- Domain 1: native heading styles ---------------------------------------
    non_native = [p.text[:40] for p in doc.paragraphs
                  if p.text.strip().startswith("#")]
    rep.check("styles.headings-native", not non_native,
              f"{len(non_native)} markdown-style headings leaked: {non_native[:3]}"
              if non_native else "no '#' heading text leaked; styles are Word-native")

    # ---- Domain 3: layout standards --------------------------------------------
    normal_size = doc.styles["Normal"].font.size
    ok_body = normal_size is None or normal_size.pt >= 10.5
    rep.check("typography.body>=10.5pt", ok_body,
              f"Normal size = {normal_size.pt if normal_size else 'inherit'}pt")
    small_runs = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.size is not None and r.font.size.pt < 9:
                            small_runs.append(r.font.size.pt)
    rep.check("typography.table>=9pt", not small_runs,
              f"{len(small_runs)} runs under 9pt" if small_runs else "all sized table runs >=9pt")
    bad_headers = []
    for t in doc.tables:
        trPr = t.rows[0]._tr.find(qn("w:trPr"))
        if trPr is None or trPr.find(qn("w:tblHeader")) is None:
            bad_headers.append(True)
    rep.check("tables.repeat-header-rows", len(doc.tables) > 0 and not bad_headers,
              f"{len(bad_headers)} tables lack repeating header row"
              if bad_headers else f"{len(doc.tables)} tables repeat headers")
    rels = [r for r in doc.part.rels.values() if "hyperlink" in r.reltype]
    rep.check("hyperlinks.native-rels", len(rels) > 0,
              f"{len(rels)} external hyperlink relationships")

    # ---- Domain 5: market-only boundary ----------------------------------------
    rules = {}
    if language_rules and Path(language_rules).exists():
        loaded = json.loads(Path(language_rules).read_text(encoding="utf-8"))
        terms = [t.lower() for t in loaded.get("prohibited_terms", [])]
    else:
        terms = DEFAULT_PROHIBITED
    survivors = [t for t in terms if t in lower]
    rep.check("boundary.prohibited-language", not survivors,
              f"survivors: {survivors}" if survivors else "no patent-analysis terminology")
    rep.check("boundary.us8969841-framing",
              "8969841" not in lower or ("inventor" in lower and "independent patent analysis" in lower),
              "US8969841 only as inventor-identified with no-analysis disclaimer")

    # ---- Domain 6: partners ------------------------------------------------------
    partner_md = run_dir / "partner-table-8530-revised.md"
    prows = _md_rows(partner_md) if partner_md.exists() else []
    orgs = [r[0] for r in prows if r]
    rep.check("partners.count>=10", len(orgs) >= 10,
              f"{len(orgs)} independent partners")
    sep_evoqua = [o for o in orgs if o.lower().startswith("evoqua")]
    rep.check("partners.evoqua-not-separate", not sep_evoqua,
              f"Evoqua listed separately from Xylem group: {sep_evoqua}" if sep_evoqua
              else "Evoqua merged into Xylem group")
    sep_noble = [o for o in orgs if o.lower().startswith("heraeus noblelight")]
    rep.check("partners.noblelight-not-separate", not sep_noble,
              f"Noblelight listed separately from Excelitas: {sep_noble}" if sep_noble
              else "Noblelight merged into Excelitas group")
    iuva_partner = [o for o in orgs if o.lower().strip().startswith(("iuva", "international ultraviolet"))]
    rep.check("partners.iuva-not-commercial-partner", not iuva_partner,
              f"IUVA appears as commercial partner: {iuva_partner}" if iuva_partner
              else "IUVA confined to associations")
    uncontacted_ok = all(
        ("not been contacted" in "|".join(r).lower()) or ("identified and verified" in "|".join(r).lower())
        for r in prows) if prows else True
    rep.check("partners.statuses-explicit", uncontacted_ok and bool(prows),
              "every partner row carries identified/verified-not-engaged status")

    # ---- Domain 6: events ---------------------------------------------------------
    events_md = run_dir / "trade-show-table-8530-revised.md"
    erows = _md_rows(events_md) if events_md.exists() else []
    etext = "\n".join("|".join(r) for r in erows)

    def _event_row(*needles):
        for r in erows:
            joined = "|".join(r).lower()
            if all(n.lower() in joined for n in needles):
                return r
        return None

    ace = _event_row("awwa annual conference", "ace27") or _event_row("ace27")
    rep.check("events.awwa-is-ace27", ace is not None,
              "AWWA entry is ACE27")
    rep.check("events.awwa-dates", bool(ace) and "june 13–16, 2027" in "|".join(ace).lower(),
              "ACE27 dated June 13–16, 2027")
    ach = _event_row("achema")
    rep.check("events.achema-dates", bool(ach) and "june 14–18, 2027" in "|".join(ach).lower(),
              "ACHEMA 2027 dated June 14–18, 2027")
    ifat = _event_row("ifat")
    rep.check("events.ifat-dates", bool(ifat) and "may 29–june 1, 2028" in "|".join(ifat).lower(),
              "IFAT Munich 2028 dated May 29–June 1, 2028")
    iuva_ev = _event_row("iuva world congress")
    iuva_txt = "|".join(iuva_ev).lower() if iuva_ev else ""
    unconfirmed_ok = ("not confirmed" in iuva_txt) and ("mid-2027" not in etext)
    rep.check("events.iuva-date-unconfirmed-stated", unconfirmed_ok,
              "IUVA shows date not confirmed; no estimated mid-2027 claim")
    combined = _event_row("windpower")
    rep.check("events.no-windpower-combined-entry", combined is None,
              "incorrect AWEA WINDPOWER/AWWA combined entry removed")

    # ---- Domain 7: source audit -----------------------------------------------------
    src_md = run_dir / "market-source-table-8530-revised.md"
    srows = _md_rows(src_md) if src_md.exists() else []
    qualifying = []
    for r in srows:
        joined = "|".join(r).lower()
        if "summary-verified" in joined or "fully verified" in joined:
            qualifying.append(r)
    rep.check("sources.qualifying>=8", len(qualifying) >= 8,
              f"{len(qualifying)} qualifying graded sources")
    audit_md = run_dir / "source-evidence-audit-8530-revised.md"
    audit_txt = audit_md.read_text(encoding="utf-8") if audit_md.exists() else ""
    sv_def = "confirmed in a publicly accessible publisher summary" in audit_txt
    rep.check("sources.summary-verified-defined", sv_def,
              "explicit summary-verified definition embedded")

    # ---- Domain 8: methodology --------------------------------------------------------
    log_md = run_dir / "research-log-8530-revised.md"
    log_txt = log_md.read_text(encoding="utf-8") if log_md.exists() else ""
    fields_needed = ["exact search string", "website / database / official-domain search used",
                     "search date", "filters or scope", "result count", "records reviewed",
                     "inclusion criteria", "exclusion criteria", "selected evidence",
                     "limitations"]
    missing_fields = [f for f in fields_needed if f not in log_txt.lower()]
    rep.check("methodology.fields-per-question", log_txt != "" and not missing_fields,
              f"missing: {missing_fields}" if missing_fields else "per-question methodology complete")
    forbidden_venues = ["google", "espacenet", "epo ops", "google scholar"]
    venue_hits = [v for v in forbidden_venues if v in log_txt.lower()]
    rep.check("methodology.no-forbidden-venues", not venue_hits,
              f"forbidden venues cited: {venue_hits}" if venue_hits
              else "official-domain research only")

    # ---- Domain 10: manifest cross-check -----------------------------------------------
    mf_path = run_dir / "run-manifest-8530-market-only-run4.json"
    if mf_path.exists():
        try:
            mf = json.loads(mf_path.read_text(encoding="utf-8"))
            rep.check("manifest.valid-json", True)
            rep.check("manifest.partner-count-matches",
                      int(mf.get("partner_count_independent", -1)) == len(orgs),
                      f"manifest={mf.get('partner_count_independent')} vs table={len(orgs)}")
        except Exception as exc:
            rep.check("manifest.valid-json", False, str(exc))
    else:
        rep.check("manifest.valid-json", False, "manifest missing")

    return rep


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument("--run-dir", required=True, type=Path)
    ap.add_argument("--language-rules", type=Path, default=None)
    args = ap.parse_args(argv)

    rep = validate(args.run_dir, args.language_rules)
    out = args.run_dir / "acceptance-report-8530-market-only-run4.json"
    out.write_text(json.dumps(rep.to_dict(), indent=2), encoding="utf-8")

    print(f"ACCEPTANCE {'PASSED' if rep.passed else 'REJECTED'}: "
          f"{len(rep.checks) - len(rep.failures)}/{len(rep.checks)} checks pass")
    for c in rep.failures:
        print(f"  FAIL {c['id']}: {c['detail']}")
    print(f"WROTE {out}")
    return 0 if rep.passed else 1


if __name__ == "__main__":
    sys.exit(main())
