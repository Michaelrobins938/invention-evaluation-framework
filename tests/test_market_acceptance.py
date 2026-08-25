"""Executable acceptance-contract tests.

Proves the validator enforces the researcher's FAIL-if list against generated
artifacts, and that a compliant full run is accepted.
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

import pytest

PKG = Path(__file__).resolve().parents[1] / "market-report-generator"
sys.path.insert(0, str(PKG))

from validate_acceptance import validate  # noqa: E402


def _generate_run(tmp_path: Path) -> Path:
    stub = tmp_path / "stub-tech.md"
    stub.write_text(
        "# Technology Overview\n\n"
        + "HPLEx is a high-power excimer source concept for water treatment. " * 8
        + "\n\nSee https://example.org for verification links.\n",
        encoding="utf-8",
    )
    run = tmp_path / "run"
    subprocess.run(
        [sys.executable, str(PKG / "generate_market_report.py"),
         "--config-dir", str(PKG / "content"), "--out", str(run),
         "--prior-tech-overview", str(stub)],
        check=True, capture_output=True,
    )
    shutil.copy2(PKG / "execution-ledger-8530-market-only-run4.md", run)
    shutil.copy2(PKG / "visual-qa-checklist-8530-run4.md", run)
    (run / "word-finalization-pending.txt").write_text("no Word in CI", encoding="utf-8")
    return run


@pytest.fixture(scope="module")
def good_run(tmp_path_factory) -> Path:
    return _generate_run(tmp_path_factory.mktemp("accept"))


def test_compliant_run_is_accepted(good_run: Path):
    rep = validate(good_run)
    assert rep.passed, [f"{c['id']}: {c['detail']}" for c in rep.failures]


def _mutate(good_run: Path, tmp_path: Path, filename: str, transform) -> Path:
    copy = tmp_path / "mutated"
    if copy.exists():
        shutil.rmtree(copy)
    shutil.copytree(good_run, copy)
    target = copy / filename
    transform(target)
    return copy


def _fail_ids(rep) -> set[str]:
    return {c["id"] for c in rep.failures}


def test_fails_on_literal_bold_markers(good_run: Path, tmp_path: Path):
    def inject(md_path: Path):
        text = md_path.read_text(encoding="utf-8")
        md_path.write_text(text.replace("# Potential Partners",
                                        "# Potential Partners\n\n**stray markdown**", 1),
                           encoding="utf-8")
        # also plant a literal ** inside the DOCX via python-docx
        from docx import Document
        docx_p = md_path.parent / "market-evaluation-8530-v17-revised.docx"
        d = Document(docx_p)
        r = d.paragraphs[1].add_run("**leaked marker**")
        r.bold = False
        d.save(docx_p)

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", inject)
    assert "markdown.no-bold-markers" in _fail_ids(validate(run))


def test_fails_when_evoqua_counted_separately(good_run: Path, tmp_path: Path):
    def inject(path: Path):
        lines = path.read_text(encoding="utf-8").splitlines()
        header_len = next(i for i, l in enumerate(lines) if l.startswith("|--"))
        rows_start = header_len + 1
        lines.insert(rows_start,
            "| Evoqua Water Technologies LLC | Independent | Sturtevant, WI, USA | Water treatment | legacy vendor | JV | identified | https://www.evoqua.com | 2026-08-25 | evidence | limits |")
        path.write_text("\n".join(lines), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", inject)
    ids = _fail_ids(validate(run))
    assert "partners.evoqua-not-separate" in ids


def test_fails_when_noblelight_counted_separately(good_run: Path, tmp_path: Path):
    def inject(path: Path):
        lines = path.read_text(encoding="utf-8").splitlines()
        header_len = next(i for i, l in enumerate(lines) if l.startswith("|--"))
        lines.insert(header_len + 1,
            "| Heraeus Noblelight GmbH | Heraeus | Hanau, Germany | UV lamps | lamp partner | supply | identified | https://www.heraeus-noblelight.com | 2026-08-25 | evidence | limits |")
        path.write_text("\n".join(lines), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", inject)
    assert "partners.noblelight-not-separate" in _fail_ids(validate(run))


def test_fails_when_iuva_listed_as_partner(good_run: Path, tmp_path: Path):
    def inject(path: Path):
        lines = path.read_text(encoding="utf-8").splitlines()
        header_len = next(i for i, l in enumerate(lines) if l.startswith("|--"))
        lines.insert(header_len + 1,
            "| IUVA Commercial Partnering Desk | IUVA | Chevy Chase, MD, USA | introductions | network | retainer | engaged | https://iuva.org | 2026-08-25 | evidence | limits |")
        path.write_text("\n".join(lines), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", inject)
    assert "partners.iuva-not-commercial-partner" in _fail_ids(validate(run))


def test_fails_below_ten_partners(good_run: Path, tmp_path: Path):
    def trim(path: Path):
        lines = path.read_text(encoding="utf-8").splitlines()
        keep = [l for l in lines]
        data_rows = [i for i, l in enumerate(keep) if l.startswith("| ") and "---" not in l][1:]
        for i in sorted(data_rows[:4], reverse=True):
            del keep[i]
        path.write_text("\n".join(keep), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", trim)
    assert "partners.count>=10" in _fail_ids(validate(run))


def test_fails_on_wrong_ace27_dates(good_run: Path, tmp_path: Path):
    def swap(path: Path):
        text = path.read_text(encoding="utf-8")
        path.write_text(text.replace("June 13–16, 2027", "June 10–13, 2027"), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "trade-show-table-8530-revised.md", swap)
    assert "events.awwa-dates" in _fail_ids(validate(run))


def test_fails_if_windpower_combined_entry_returns(good_run: Path, tmp_path: Path):
    def inject(path: Path):
        lines = path.read_text(encoding="utf-8").splitlines()
        header_len = next(i for i, l in enumerate(lines) if l.startswith("|--"))
        lines.insert(header_len + 1,
            "| AWEA WINDPOWER / AWWA Annual Conference | ACP/AWWA | mixed | TBD | annual | wind+water | attend | https://example.org | 2026-08-25 | v | v |")
        path.write_text("\n".join(lines), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "trade-show-table-8530-revised.md", inject)
    assert "events.no-windpower-combined-entry" in _fail_ids(validate(run))


def test_fails_if_iuva_date_presented_as_confirmed(good_run: Path, tmp_path: Path):
    def swap(path: Path):
        text = path.read_text(encoding="utf-8")
        path.write_text(text.replace("Date not confirmed", "Confirmed mid-2027"), encoding="utf-8")

    run = _mutate(good_run, tmp_path, "trade-show-table-8530-revised.md", swap)
    assert "events.iuva-date-unconfirmed-stated" in _fail_ids(validate(run))


def test_fails_below_eight_qualifying_sources(good_run: Path, tmp_path: Path):
    def degrade(path: Path):
        text = path.read_text(encoding="utf-8")
        for grade in ("summary-verified", "fully verified"):
            text = text.replace(grade, "unretrieved citation")
        path.write_text(text, encoding="utf-8")

    run = _mutate(good_run, tmp_path, "market-source-table-8530-revised.md", degrade)
    assert "sources.qualifying>=8" in _fail_ids(validate(run))


def test_fails_on_prohibited_patent_language(good_run: Path, tmp_path: Path):
    def inject(path: Path):
        from docx import Document
        p = path.parent / "market-evaluation-8530-v17-revised.docx"
        d = Document(p)
        d.paragraphs[1].add_run(" The closest prior art anticipates claim 1.")
        d.save(p)

    run = _mutate(good_run, tmp_path, "partner-table-8530-revised.md", inject)
    ids = _fail_ids(validate(run))
    assert "boundary.prohibited-language" in ids


def test_acceptance_report_written_by_cli(good_run: Path, tmp_path: Path):
    run = tmp_path / "cli-run"
    shutil.copytree(good_run, run)
    proc = subprocess.run(
        [sys.executable, str(PKG / "validate_acceptance.py"), "--run-dir", str(run)],
        capture_output=True, text=True,
    )
    assert proc.returncode == 0
    report = json.loads((run / "acceptance-report-8530-market-only-run4.json").read_text())
    assert report["accepted"] is True
    assert report["total_checks"] >= 25
